"""DOCX parser using python-docx — paragraph-level extraction grouped into logical pages."""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from backend.core.exceptions import IngestError
from backend.ingestion.parsers.base import DocumentParser, ParsedDocument, ParsedPage

# Group every N paragraphs into a "page" for citation purposes
_PARAS_PER_PAGE = 20


class DOCXParser(DocumentParser):
    """Parse DOCX files. Groups paragraphs into logical pages for citation tracking."""

    def supports(self, file_path: str | Path) -> bool:
        return Path(file_path).suffix.lower() in {".docx", ".doc"}

    def parse(self, file_path: str | Path, title: str | None = None) -> ParsedDocument:
        path = Path(file_path)
        if not path.exists():
            raise IngestError(f"File not found: {file_path}")
        if not self.supports(path):
            raise IngestError(f"DOCXParser does not support file type: {path.suffix}")

        display_title = title or path.stem

        try:
            doc = DocxDocument(str(path))
        except Exception as exc:
            raise IngestError(f"Failed to open DOCX {file_path}: {exc}") from exc

        all_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        if not all_paragraphs:
            return ParsedDocument(source_path=str(path), title=display_title, pages=[])

        # Group paragraphs into logical pages
        pages = []
        for page_idx, start in enumerate(range(0, len(all_paragraphs), _PARAS_PER_PAGE), start=1):
            chunk_paras = all_paragraphs[start : start + _PARAS_PER_PAGE]
            page_text = "\n\n".join(chunk_paras)
            pages.append(
                ParsedPage(
                    page_number=page_idx,
                    text=page_text,
                    paragraphs=chunk_paras,
                )
            )

        return ParsedDocument(source_path=str(path), title=display_title, pages=pages)
